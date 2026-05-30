#!/usr/bin/env python3
from __future__ import annotations

import argparse
import datetime as dt
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path


DEFAULT_VAULT = Path("/Users/zydeng/Library/Mobile Documents/iCloud~md~obsidian/Documents/Purple Sys")
TEMP_ROOT = Path("/private/tmp/purple-doc-bridge")
PYTHON_DOCX_AVAILABLE = True
try:
    import docx
except Exception:
    PYTHON_DOCX_AVAILABLE = False
    docx = None  # type: ignore[assignment]


def run(cmd: list[str], *, capture: bool = False) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        cmd,
        check=True,
        text=True,
        stdout=subprocess.PIPE if capture else None,
        stderr=subprocess.PIPE if capture else None,
    )


def extract_with_pandoc(path: Path) -> str:
    with tempfile.TemporaryDirectory(prefix="import-", dir=TEMP_ROOT) as raw_tmp:
        tmp = Path(raw_tmp)
        extracted = tmp / "extracted.md"
        run(["pandoc", str(path), "-f", "docx", "-t", "markdown", "--wrap=none", "-o", str(extracted)])
        return extracted.read_text(encoding="utf-8").strip()


def extract_with_python_docx(path: Path) -> str:
    if not PYTHON_DOCX_AVAILABLE or docx is None:
        raise RuntimeError("python-docx is not available for fallback extraction")
    document = docx.Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
        if rows:
            width = max(len(row) for row in rows)
            normalized = [row + [""] * (width - len(row)) for row in rows]
            parts.append("")
            parts.append("| " + " | ".join(normalized[0]) + " |")
            parts.append("| " + " | ".join(["---"] * width) + " |")
            for row in normalized[1:]:
                parts.append("| " + " | ".join(row) + " |")
    return "\n\n".join(parts).strip()


def extract_docx(path: Path) -> str:
    TEMP_ROOT.mkdir(parents=True, exist_ok=True)
    try:
        return extract_with_pandoc(path)
    except Exception:
        return extract_with_python_docx(path)


def normalize_chinese_quotes(text: str) -> str:
    text = text.replace(r"\"", '"')
    output: list[str] = []
    in_code_block = False
    quote_pattern = re.compile(r'"([^"\n]*[\u4e00-\u9fff][^"\n]*)"')
    for line in text.splitlines():
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            output.append(line)
            continue
        if in_code_block:
            output.append(line)
            continue
        output.append(quote_pattern.sub(lambda match: f"“{match.group(1)}”", line))
    return "\n".join(output)


def normalize_markdown_structure(text: str, *, title: str) -> str:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").splitlines()
    normalized: list[str] = []

    for raw_line in lines:
        line = raw_line.replace("\t", "   ").rstrip()
        stripped = line.strip()

        if not stripped:
            normalized.append("")
            continue

        section = re.fullmatch(r"\*\*([一二三四五六七八九十]+、[^*]+)\*\*", stripped)
        if section:
            normalized.append(f"## {section.group(1).strip()}")
            continue

        subsection = re.fullmatch(r"（([一二三四五六七八九十]+)）(.+)", stripped)
        if subsection:
            normalized.append(f"### （{subsection.group(1)}）{subsection.group(2).strip()}")
            continue

        numbered = re.match(r"^(\s*)(\d+)(?:[、．]|\\\.)\s*(.+)$", line)
        if numbered:
            normalized.append(f"{numbered.group(1)}{numbered.group(2)}. {numbered.group(3).strip()}")
            continue

        normalized.append(line)

    normalized = normalize_list_spacing(normalized)
    text = "\n".join(normalized).strip()
    text = normalize_chinese_quotes(text)
    return remove_duplicate_opening_title(text, title=title)


def is_ordered_list_item(line: str) -> bool:
    return bool(re.match(r"^\s*\d+\.\s+", line))


def is_heading_or_list(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("#") or is_ordered_list_item(line) or stripped.startswith("- ")


def list_item_needs_continuation(line: str) -> bool:
    stripped = line.strip()
    if not is_ordered_list_item(stripped):
        return False
    content = re.sub(r"^\d+\.\s+", "", stripped)
    return len(content) <= 24 and not content.endswith(("。", "；", "，", "：", ".", ";", ",", ":"))


def previous_nonempty(lines: list[str], index: int) -> str:
    for cursor in range(index - 1, -1, -1):
        if lines[cursor].strip():
            return lines[cursor]
    return ""


def next_nonempty(lines: list[str], index: int) -> str:
    for cursor in range(index + 1, len(lines)):
        if lines[cursor].strip():
            return lines[cursor]
    return ""


def normalize_list_spacing(lines: list[str]) -> list[str]:
    compacted: list[str] = []
    for index, line in enumerate(lines):
        if line.strip():
            prev = previous_nonempty(lines, index)
            if list_item_needs_continuation(prev) and not is_heading_or_list(line):
                compacted.append(f"   {line.strip()}")
            else:
                compacted.append(line)
            continue

        prev = previous_nonempty(lines, index)
        nxt = next_nonempty(lines, index)
        if is_ordered_list_item(prev) and is_ordered_list_item(nxt):
            continue
        compacted.append(line)
    return compacted


def remove_duplicate_opening_title(text: str, *, title: str) -> str:
    lines = text.splitlines()
    first_index = next((index for index, line in enumerate(lines) if line.strip()), None)
    if first_index is None:
        return text

    first = lines[first_index].strip().lstrip("#").strip()
    normalized_title = title
    if normalized_title[:8].isdigit() and len(normalized_title) > 9:
        normalized_title = normalized_title[9:].strip()
    normalized_title = re.sub(r"（[^）]*稿）$", "", normalized_title).strip()

    if first == title or (normalized_title and first == normalized_title):
        del lines[first_index]
        return "\n".join(lines).strip()
    return text


def yaml_scalar(value: str) -> str:
    if value == "":
        return ""
    if any(ch in value for ch in [":", "#", "[", "]", "{", "}", ",", "\"", "'"]):
        return '"' + value.replace('"', '\\"') + '"'
    return value


def frontmatter(*, material_type: str, status: str, project: str, created: str) -> str:
    return "\n".join(
        [
            "---",
            "type: material",
            f"material_type: {yaml_scalar(material_type)}",
            f"status: {yaml_scalar(status)}",
            f"project: {yaml_scalar(project)}",
            f"created: {yaml_scalar(created)}",
            "---",
            "",
        ]
    )


def safe_filename(title: str) -> str:
    return title.replace("/", "-").replace(":", "：").strip()


def resolve_output(vault: Path, title: str, output: str | None) -> Path:
    if output:
        raw = Path(output).expanduser()
        return raw if raw.is_absolute() else vault / raw
    return vault / "03 Material" / f"{safe_filename(title)}.md"


def write_material(
    docx_path: Path,
    vault: Path,
    *,
    title: str,
    output: str | None,
    material_type: str,
    status: str,
    project: str,
    created: str,
    attach_original: bool,
    daily_note: str | None,
) -> Path:
    if not docx_path.exists():
        raise FileNotFoundError(docx_path)

    conventions = vault / "SYSTEM-CONVENTIONS.md"
    template = vault / "99 Template" / "Material" / f"{material_type.capitalize()} Template.md"
    note_template = vault / "99 Template" / "Material" / "Note Template.md"
    if not conventions.exists():
        raise FileNotFoundError(f"Purple Sys conventions not found: {conventions}")
    if not template.exists() and not note_template.exists():
        raise FileNotFoundError("No Purple Sys Material template found")

    output_path = resolve_output(vault, title, output).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    body = normalize_markdown_structure(extract_docx(docx_path), title=title)
    attachment_line = ""
    if attach_original:
        attachment_dir = vault / "03 Material" / "attachments"
        attachment_dir.mkdir(parents=True, exist_ok=True)
        attachment_path = attachment_dir / docx_path.name
        shutil.copy2(docx_path, attachment_path)
        attachment_line = f"\n## 来源文件\n\n- [[attachments/{attachment_path.name}]]\n\n"

    text = frontmatter(
        material_type=material_type,
        status=status,
        project=project,
        created=created,
    )
    text += f"# {title}\n"
    text += attachment_line
    if body:
        text += "\n" + body.strip() + "\n"
    output_path.write_text(text, encoding="utf-8")

    if daily_note:
        daily_path = Path(daily_note).expanduser()
        if not daily_path.is_absolute():
            daily_path = vault / daily_path
        daily_path.parent.mkdir(parents=True, exist_ok=True)
        with daily_path.open("a", encoding="utf-8") as handle:
            handle.write(f"\n- 新增 Material：[[{output_path.stem}]]\n")

    return output_path


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Import a DOCX into Purple Sys 03 Material.")
    parser.add_argument("--docx", required=True, help="Source DOCX path.")
    parser.add_argument("--title", required=True, help="Material note title.")
    parser.add_argument("--vault", default=str(DEFAULT_VAULT), help="Purple Sys vault path.")
    parser.add_argument("--output", help="Output markdown path, absolute or relative to vault.")
    parser.add_argument("--material-type", default="note", help="Purple Sys material_type value.")
    parser.add_argument("--status", default="inbox", help="Purple Sys status value.")
    parser.add_argument("--project", default="", help="Purple Sys project value.")
    parser.add_argument("--created", default=dt.date.today().isoformat(), help="Created date.")
    parser.add_argument("--attach-original", action="store_true", help="Copy original DOCX to 03 Material/attachments.")
    parser.add_argument("--daily-note", help="Optional daily note path to append a short backlink.")
    return parser.parse_args(argv)


def main(argv: list[str]) -> int:
    args = parse_args(argv)
    try:
        output = write_material(
            Path(args.docx).expanduser(),
            Path(args.vault).expanduser(),
            title=args.title,
            output=args.output,
            material_type=args.material_type,
            status=args.status,
            project=args.project,
            created=args.created,
            attach_original=args.attach_original,
            daily_note=args.daily_note,
        )
    except Exception as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 1

    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
